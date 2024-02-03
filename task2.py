# relies on the usage of python-docx library
# as I didn't encounter writing docx files programmatically before, I looked into them more
# resulting from that, I believe letting a library is better, as it provides more flexibility, rather than 
# having to write and edit the xml structure manually and write into the file (as the library provides more comfort and safety)
#
# of course, writing some more advanced handling for it is possible, though it seemed overkill for a test task

from docx import Document 

def add_row(table, key, value):
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value

def create_docx(name, id_number, sampling_date):
    document = Document()

    heading = document.add_heading()
    heading.add_run("Výsledný protokol genetického vyšetření").bold = True
    heading.alignment = 1 # as per documentation, 1 is centered

    table = document.add_table(rows=0, cols=2)
    add_row(table, "Jméno a příjmení:", name)
    add_row(table, "Rodné číslo:", id_number)
    add_row(table, "Datum odběru:", sampling_date)

    document.add_page_break()
    document.save('result.docx')


create_docx("Alice Bob", "12345/6789", "2021/2/2")